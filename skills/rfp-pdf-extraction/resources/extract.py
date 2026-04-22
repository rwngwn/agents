# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "pdfplumber>=0.11",
#   "python-docx>=1.1",
#   "pytesseract>=0.3.10",
#   "pdf2image>=1.17",
#   "Pillow>=10.0",
# ]
# ///
"""
RFP document extractor — PDF, DOCX, TXT, MD → clean UTF-8 markdown.

Run via: uv run --script extract.py --input file.pdf --output out.md

Features:
- pdfplumber primary, pdftotext fallback, tesseract OCR for scans
- NFC Unicode normalization (Czech diacritics)
- Header/footer stripping (repeating lines)
- Table extraction as markdown tables
- Page markers: ## Page N
"""
from __future__ import annotations
import argparse
import collections
import re
import shutil
import subprocess
import sys
import unicodedata
from pathlib import Path


def normalize(text: str) -> str:
    if not text:
        return ""
    return unicodedata.normalize("NFC", text)


def strip_headers_footers(pages: list[str]) -> list[str]:
    """Remove lines that repeat on >= 50% of pages (likely headers/footers)."""
    if len(pages) < 3:
        return pages
    line_counts: collections.Counter = collections.Counter()
    page_line_sets = []
    for p in pages:
        lines = {ln.strip() for ln in p.splitlines() if ln.strip()}
        page_line_sets.append(lines)
        line_counts.update(lines)
    threshold = max(2, len(pages) // 2)
    repeating = {ln for ln, c in line_counts.items() if c >= threshold and len(ln) < 120}
    cleaned = []
    for p in pages:
        out_lines = [ln for ln in p.splitlines() if ln.strip() not in repeating]
        cleaned.append("\n".join(out_lines))
    return cleaned


def table_to_md(table: list[list[str | None]]) -> str:
    if not table or not table[0]:
        return ""
    header = [str(c or "").strip().replace("\n", " ") for c in table[0]]
    rows = [[str(c or "").strip().replace("\n", " ") for c in r] for r in table[1:]]
    cols = len(header)
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(["---"] * cols) + " |")
    for r in rows:
        r = r + [""] * (cols - len(r))
        lines.append("| " + " | ".join(r[:cols]) + " |")
    return "\n".join(lines)


def extract_pdf_pdfplumber(path: Path) -> list[dict]:
    import pdfplumber  # type: ignore

    out = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = normalize(page.extract_text() or "")
            tables = []
            try:
                for t in page.extract_tables() or []:
                    md = table_to_md(t)
                    if md:
                        tables.append(md)
            except Exception:
                pass
            out.append({"page": i, "text": text, "tables": tables})
    return out


def extract_pdf_pdftotext(path: Path) -> list[dict]:
    """Fallback using poppler's pdftotext -layout."""
    if not shutil.which("pdftotext"):
        return []
    try:
        raw = subprocess.check_output(
            ["pdftotext", "-layout", "-enc", "UTF-8", str(path), "-"],
            text=True,
            encoding="utf-8",
            stderr=subprocess.DEVNULL,
        )
    except subprocess.CalledProcessError:
        return []
    # pdftotext separates pages with \f (form feed)
    raw = normalize(raw)
    pages = raw.split("\f")
    return [{"page": i, "text": p.strip(), "tables": []} for i, p in enumerate(pages, 1) if p.strip()]


def extract_pdf_ocr(path: Path) -> list[dict]:
    try:
        from pdf2image import convert_from_path  # type: ignore
        import pytesseract  # type: ignore
    except ImportError:
        print("[extract] OCR libs not available", file=sys.stderr)
        return []
    if not shutil.which("tesseract"):
        print("[extract] tesseract not in PATH — install with: brew install tesseract tesseract-lang", file=sys.stderr)
        return []
    images = convert_from_path(str(path), dpi=300)
    out = []
    for i, img in enumerate(images, 1):
        text = normalize(pytesseract.image_to_string(img, lang="ces+eng"))
        out.append({"page": i, "text": text, "tables": []})
    return out


def extract_docx(path: Path) -> list[dict]:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        print("[extract] python-docx not available", file=sys.stderr)
        return []
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        txt = normalize(p.text)
        if txt:
            parts.append(txt)
    tables = []
    for t in doc.tables:
        rows = []
        for row in t.rows:
            rows.append([normalize(cell.text).replace("\n", " ") for cell in row.cells])
        if rows:
            md = table_to_md(rows)
            if md:
                tables.append(md)
    # DOCX has no real pages — treat whole doc as 1 page
    return [{"page": 1, "text": "\n\n".join(parts), "tables": tables}]


def extract_text(path: Path) -> list[dict]:
    content = normalize(path.read_text(encoding="utf-8", errors="replace"))
    return [{"page": 1, "text": content, "tables": []}]


def is_pdf_text_empty(pages: list[dict], min_chars: int = 200) -> bool:
    total = sum(len(p["text"]) for p in pages)
    return total < min_chars


def render_output(pages: list[dict], source: Path) -> str:
    lines = [f"<!-- Extracted from: {source} -->", ""]
    for p in pages:
        if not p["text"] and not p["tables"]:
            continue
        lines.append(f"## Page {p['page']}")
        lines.append("")
        if p["text"]:
            lines.append(p["text"].strip())
            lines.append("")
        for i, t in enumerate(p["tables"], 1):
            lines.append(f"**Table {p['page']}.{i}**")
            lines.append("")
            lines.append(t)
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True, type=Path)
    ap.add_argument("--output", required=True, type=Path)
    ap.add_argument("--ocr", action="store_true", help="Force OCR (for scanned PDFs)")
    ap.add_argument("--normalize", default="nfc", choices=["nfc", "nfkc", "none"])
    args = ap.parse_args()

    src: Path = args.input
    if not src.exists():
        print(f"[extract] Input not found: {src}", file=sys.stderr)
        return 2

    ext = src.suffix.lower()
    pages: list[dict] = []

    if ext == ".pdf":
        if args.ocr:
            pages = extract_pdf_ocr(src)
        else:
            pages = extract_pdf_pdfplumber(src)
            if is_pdf_text_empty(pages):
                print("[extract] pdfplumber yielded <200 chars — trying pdftotext fallback", file=sys.stderr)
                alt = extract_pdf_pdftotext(src)
                if not is_pdf_text_empty(alt):
                    pages = alt
                else:
                    print("[extract] pdftotext also empty — falling back to OCR", file=sys.stderr)
                    pages = extract_pdf_ocr(src)
    elif ext in (".docx",):
        pages = extract_docx(src)
    elif ext in (".txt", ".md"):
        pages = extract_text(src)
    else:
        print(f"[extract] Unsupported extension: {ext}", file=sys.stderr)
        return 2

    if not pages or is_pdf_text_empty(pages, min_chars=100):
        print(f"[extract] No usable text extracted from {src}", file=sys.stderr)
        return 3

    # Strip repeating headers/footers for multi-page PDFs
    if ext == ".pdf" and len(pages) >= 3:
        texts = [p["text"] for p in pages]
        cleaned = strip_headers_footers(texts)
        for p, t in zip(pages, cleaned):
            p["text"] = t

    output = render_output(pages, src)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(output, encoding="utf-8")

    total_chars = sum(len(p["text"]) for p in pages)
    total_tables = sum(len(p["tables"]) for p in pages)
    print(
        f"[extract] OK — {len(pages)} pages, {total_chars} chars, "
        f"{total_tables} tables → {args.output}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
